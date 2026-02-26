"""
Module đọc và tiền xử lý file .docx
"""

from docx import Document
import re

class DocumentReader:
    """
    Class đọc file .docx và trích xuất nội dung văn bản
    """
    
    def __init__(self, file_path):
        """
        Khởi tạo DocumentReader
        
        Args:
            file_path (str): Đường dẫn tới file .docx
        """
        self.file_path = file_path
        self.document = None
        self.paragraphs = []
    
    def load(self):
        """
        Đọc file .docx và lưu nội dung
        
        Returns:
            bool: True nếu đọc thành công
        """
        try:
            self.document = Document(self.file_path)
            return True
        except Exception as e:
            print(f"Lỗi khi đọc file: {e}")
            return False
    
    def extract_paragraphs(self):
        """
        Trích xuất các đoạn văn từ document
        Bỏ qua các dòng trống và header/footer
        
        Returns:
            list: Danh sách các đoạn văn (paragraphs)
        """
        if not self.document:
            return []
        
        paragraphs = []
        for para in self.document.paragraphs:
            text = para.text.strip()
            
            # Bỏ qua dòng trống
            if not text:
                continue
            
            # Bỏ qua các dòng chỉ chứa số trang hoặc header đơn giản
            if self._is_header_footer(text):
                continue
            
            paragraphs.append(text)
        
        self.paragraphs = paragraphs
        return paragraphs
    
    def _is_header_footer(self, text):
        """
        Kiểm tra xem text có phải là header/footer không
        
        Args:
            text (str): Text cần kiểm tra
            
        Returns:
            bool: True nếu là header/footer
        """
        # Bỏ qua các dòng chỉ chứa số
        if re.match(r'^\d+$', text):
            return True
        
        # Bỏ qua các dòng kiểu "Trang 1", "Page 1"
        if re.match(r'^(Trang|Page)\s+\d+', text, re.IGNORECASE):
            return True
        
        # Bỏ qua các dòng quá ngắn (< 3 ký tự) trừ khi là số điều, khoản
        if len(text) < 3 and not re.match(r'^\d+\.$', text):
            return True
        
        return False
    
    def get_paragraphs(self):
        """
        Lấy danh sách paragraphs đã được xử lý
        
        Returns:
            list: Danh sách paragraphs
        """
        return self.paragraphs
    
    def read(self):
        """
        Đọc file và trả về danh sách paragraphs
        
        Returns:
            list: Danh sách paragraphs hoặc None nếu lỗi
        """
        if self.load():
            return self.extract_paragraphs()
        return None
